"""
DOCX Yazıcı Modülü v2
Şablon DOCX dosyasını verilerle doldurur.
Güncellemeler:
- Ana Dağıtım Pano bilgileri desteği
- Tüm kontrol kriterleri doldurma
- Calibri 6pt font (Linye Adı sütunu)
- Otomatik kusur açıklaması
- Iz otomatik hesaplama
- 6.2 Potansiyel Dengeleme tablosu sayfa sonu
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Any, Optional
import os

# Import yapılandırması - Modül olarak veya script olarak çalışmaya uyum
try:
    import constants as const
except ImportError:
    from . import constants as const

class DocxWriter:
    """DOCX şablon dosyasını verilerle doldurur."""

    def __init__(self, template_path: str):
        """
        Args:
            template_path: Şablon DOCX dosyasının yolu
        """
        self.template_path = template_path
        self.doc = None

    def load_template(self) -> bool:
        """Şablon dosyasını yükler."""
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Sablon bulunamadi: {self.template_path}")

        self.doc = Document(self.template_path)
        return True

    def _copy_row(self, table, source_row_idx: int):
        """Bir satırı kopyalayarak tabloya yeni satır ekler (format korunur).

        Args:
            table: Tablo objesi
            source_row_idx: Kopyalanacak kaynak satır indeksi

        Returns:
            Yeni eklenen satır objesi
        """
        from copy import deepcopy
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Kaynak satırı al
        source_row = table.rows[source_row_idx]

        # Satır XML'ini kopyala
        new_tr = deepcopy(source_row._tr)

        # Tabloya ekle
        table._tbl.append(new_tr)

        # Yeni satırı al (son satır)
        new_row = table.rows[-1]

        # Tüm hücrelerin içeriğini temizle ama formatı koru
        for cell in new_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""
                # Paragrafın direkt textini de temizle
                if para.text:
                    para.clear()

        return new_row

    def _get_cell_safe(self, table, row_idx: int, col_idx: int):
        """Güvenli hücre erişimi."""
        try:
            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                return table.rows[row_idx].cells[col_idx]
        except (IndexError, AttributeError):
            pass
        return None

    def _replace_placeholder_in_cell(self, cell, placeholder: str, value: str, font_size: int = 9):
        """Hücredeki placeholder'ı değerle değiştirir, font boyutu 9pt.
        Word run-splitting sorununu aşmak için paragraf metnini değiştirir.
        """
        if cell is None or not placeholder or value is None:
            return

        # Önce basit run taraması (performans ve stil koruma için)
        replaced = False
        for para in cell.paragraphs:
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
                    run.font.size = Pt(font_size)
                    replaced = True

            # Eğer run içinde bulunamadıysa ama paragrafta varsa (split run durumu)
            if not replaced and placeholder in para.text:
                # Paragraf metnini komple değiştir (stilleri sıfırlar ama çalışır)
                para.text = para.text.replace(placeholder, str(value))
                # Yeni oluşan run için font ayarla
                for run in para.runs:
                    run.font.size = Pt(font_size)
                replaced = True

    def _replace_placeholder_with_format(self, cell, placeholder: str, value: str,
                                          font_name: str = 'Calibri', font_size: float = 7, bold: bool = True):
        """Hücredeki placeholder'ı özel formatla değiştirir (Calibri Kalın 7pt gibi)."""
        if cell is None or not placeholder or value is None:
            return

        for para in cell.paragraphs:
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.bold = bold
                    return

            # Split run durumu
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.bold = bold
                return

    def _set_cell_value(self, cell, value: str, font_name: str = None, font_size: int = 9):
        """Hücre değerini ayarlar, mevcut içeriği TAMAMEN temizler. Varsayılan font 9pt."""
        if cell is None or value is None:
            return

        from docx.shared import Pt

        # TÜM paragrafları temizle (birden fazla paragraf olabilir)
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""
            # Paragrafın kendi text'ini de temizle
            if hasattr(para, '_p'):
                for child in list(para._p):
                    if child.tag.endswith('}r') or child.tag.endswith('}bookmarkStart') or child.tag.endswith('}bookmarkEnd'):
                        continue

        # İlk paragrafa yeni değeri yaz
        if cell.paragraphs:
            para = cell.paragraphs[0]
            # Tüm run'ları temizle
            for run in para.runs:
                run.text = ""

            # Paragraf boşluklarını sıfırla
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.0  # Tek satır aralığı

            # İlk run'a yaz veya yeni run ekle
            if para.runs:
                run = para.runs[0]
                run.text = str(value)
            else:
                run = para.add_run(str(value))

            # Font ayarları - varsayılan 9pt
            if font_name:
                run.font.name = font_name
            run.font.size = Pt(font_size)  # Her zaman font boyutu uygula
        else:
            cell.text = str(value)

    def fill_firma_bilgileri(self, data: Dict[str, Any]):
        """Firma bilgilerini doldurur (Tablo 1).
        Kullanıcının belirlediği placeholder'ları değiştirir.

        Placeholderlar:
        - firma_adi
        - kontrol_adresi
        - sgk_sicil
        - rapor_numarasi
        - rapor_tarihi
        - sozlesme_id
        - baslangic_tarih_saat
        - bitis_tarih_saat
        - bir_sonraki_kontrol
        """
        if not self.doc or not self.doc.tables:
            return

        table = self.doc.tables[0]  # İlk tablo

        # Tarih işlemleri
        rapor_tarihi = data.get('Rapor Tarihi', '')
        # GUI'den gelen başlangıç/bitiş tarihleri
        baslangic = data.get('Kontrol Baslangic', '')
        bitis = data.get('Kontrol Bitis', '')
        sonraki_kontrol = ''

        if rapor_tarihi:
            try:
                # Tarih formatı dd.mm.yyyy varsayılıyor
                parts = rapor_tarihi.split('.')
                if len(parts) == 3:
                    import datetime
                    dt = datetime.datetime(int(parts[2]), int(parts[1]), int(parts[0]))

                    # Sonraki kontrol: 1 yıl sonra
                    next_year = dt.replace(year=dt.year + 1)
                    sonraki_kontrol = next_year.strftime("%d.%m.%Y")

                    # Eğer başlangıç/bitiş girilmediyse varsayılan oluştur
                    if not baslangic:
                        baslangic = f"{rapor_tarihi} 09:00"
                    if not bitis:
                        bitis = f"{rapor_tarihi} 17:00"
            except Exception:
                pass

        # Placeholder -> Değer eşleştirmesi
        placeholders = {
            'firma_adi': data.get('Firma Adi', ''),
            'kontrol_adresi': data.get('Tesis Adresi', ''),
            'sgk_sicil': data.get('SGK Sicil No', ''),
            'rapor_numarasi': data.get('Rapor No', ''),
            'rapor_tarihi': rapor_tarihi,
            'sozlesme_id': data.get('Sozlesme ID', ''),
            'baslangic_tarih_saat': baslangic,
            'bitis_tarih_saat': bitis,
            'bir_sonraki_kontrol': sonraki_kontrol,
            'tklf': data.get('Teklif Numarasi', ''),
            'isim_soyisim': data.get('Kontrol Eden', ''),
            'belge_no': data.get('Belge No', ''),
            # Termal Kamera placeholder'ları
            'termal_cihaz_adi': data.get('Termal Cihaz Adi', ''),
            'termal_kalibrasyon_tarihi': data.get('Termal Kalibrasyon Tarihi', ''),
            'termal_kalibrasyon_gecerlilik': data.get('Termal Kalibrasyon Gecerlilik', ''),
            'termal_seri_numarasi': data.get('Termal Seri No', ''),
            'termal_kalibrasyon_no': data.get('Termal Kalibrasyon No', ''),
            # Ölçüm Cihazı placeholder'ları
            'olcum_cihaz_adi': data.get('Olcum Cihaz Adi', ''),
            'olcum_kalibrasyon_tarihi': data.get('Olcum Kalibrasyon Tarihi', ''),
            'olcum_kalibrasyon_gecerlilik': data.get('Olcum Kalibrasyon Gecerlilik', ''),
            'olcum_seri_numarasi': data.get('Olcum Seri No', ''),
            'olcum_kalibrasyon_no': data.get('Olcum Kalibrasyon No', ''),
        }

        # Tablo 1'deki hücreleri gez ve placeholderları değiştir
        for row in table.rows:
            for cell in row.cells:
                # Her placeholder için kontrol et
                for ph, value in placeholders.items():
                    if ph in cell.text and value:
                        self._replace_placeholder_in_cell(cell, ph, str(value))

        # TÜM tablolarda global placeholder'ları ara
        # (11. PERİYODİK KONTROL YETKİLİ KİŞİ, ONAY, CİHAZ BİLGİLERİ bölümleri için)
        global_placeholders = [
            'isim_soyisim', 'belge_no', 'tklf',
            'termal_cihaz_adi', 'termal_kalibrasyon_tarihi', 'termal_kalibrasyon_gecerlilik',
            'termal_seri_numarasi', 'termal_kalibrasyon_no',
            'olcum_cihaz_adi', 'olcum_kalibrasyon_tarihi', 'olcum_kalibrasyon_gecerlilik',
            'olcum_seri_numarasi', 'olcum_kalibrasyon_no'
        ]
        for tbl in self.doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for ph in global_placeholders:
                        if ph in cell.text and placeholders.get(ph):
                            # isim_soyisim ve belge_no için Calibri Kalın 7pt
                            if ph in ['isim_soyisim', 'belge_no']:
                                self._replace_placeholder_with_format(cell, ph, str(placeholders[ph]),
                                                                       font_name='Calibri', font_size=7, bold=True)
                            else:
                                self._replace_placeholder_in_cell(cell, ph, str(placeholders[ph]))

        # Paragraflar içinde de ara (tablo dışındaki alanlar için)
        for para in self.doc.paragraphs:
            for ph in global_placeholders:
                if ph in para.text and placeholders.get(ph):
                    for run in para.runs:
                        if ph in run.text:
                            run.text = run.text.replace(ph, str(placeholders[ph]))
                            # isim_soyisim ve belge_no için Calibri Kalın 7pt
                            if ph in ['isim_soyisim', 'belge_no']:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(7)
                                run.font.bold = True

        # === PROJE VE TEK HAT ŞEMASI CHECKBOX'LARI ===
        # Satır 11'deki checkbox'ları işle
        proje_var = data.get('Proje Var', 'Var')
        tekhat_var = data.get('Tek Hat Var', 'Var')

        print(f"[INFO] Proje Var: {proje_var}, Tek Hat Var: {tekhat_var}")

        if len(table.rows) > 11:
            row = table.rows[11]

            # Proje Var mı? checkbox'ları (Hücre 7-8)
            for col_idx in [7, 8]:
                cell = self._get_cell_safe(table, 11, col_idx)
                if cell:
                    cell_text = cell.text
                    if 'Var' in cell_text or 'Yok' in cell_text:
                        # Temiz format: seçilen değere göre işaretle
                        if proje_var == "Var":
                            new_text = "Var: [X]\tYok: [ ]"
                        else:
                            new_text = "Var: [ ]\tYok: [X]"
                        self._set_cell_value(cell, new_text, font_size=9)
                        break  # Sadece bir kez yaz (merged hücreler)

            # Tek Hat Şeması Var mı? checkbox'ları (Hücre 12)
            cell = self._get_cell_safe(table, 11, 12)
            if cell:
                cell_text = cell.text
                if 'Var' in cell_text or 'Yok' in cell_text:
                    if tekhat_var == "Var":
                        new_text = "Var: [X]\tYok: [ ]"
                    else:
                        new_text = "Var: [ ]\tYok: [X]"
                    self._set_cell_value(cell, new_text, font_size=9)

        # === YAPI CİNSİ CHECKBOX'LARI ===
        # Satır 13'te Yapı Cinsi seçenekleri var (Ev, Ticari, Endüstri, Diğer)
        yapi_cinsi = data.get('Yapi Cinsi', 'Ticari')
        print(f"[INFO] Yapı Cinsi: {yapi_cinsi}")

        if len(table.rows) > 13:
            cell = self._get_cell_safe(table, 13, 1)
            if cell:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

                # Hücre dikey hizalamasını TOP yap
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

                # Hücredeki TÜM paragrafları XML seviyesinde sil (ilki hariç)
                tc = cell._tc
                # Tüm <w:p> elementlerini bul
                p_elements = tc.findall(qn('w:p'))

                # İlk paragraf hariç hepsini sil
                for p_elem in p_elements[1:]:
                    tc.remove(p_elem)

                # İlk paragrafın içeriğini temizle
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    # Paragraftaki tüm run'ları XML seviyesinde sil
                    for child in list(para._p):
                        if child.tag.endswith('}r'):
                            para._p.remove(child)

                    # Paragraf formatını ayarla
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.line_spacing = 1.0

                    # İçeriği yaz
                    options = ["Ev", "Ticari", "Endüstri", "Diğer"]
                    lines = []
                    for option in options:
                        if option == yapi_cinsi:
                            lines.append(f"[X]{option}")
                        else:
                            lines.append(f"[ ]{option}")

                    run = para.add_run("\n".join(lines))
                    run.font.size = Pt(9)

    def fill_potansiyel_dengeleme_ve_zemin(self, data: Dict[str, Any], fonksiyon_testleri: List[Dict[str, Any]] = None):
        """6.2 Potansiyel Dengeleme ve 6.3 Zemin İzolasyonu bölümlerini doldurur.

        Tablo 3 (index 3):
        - 6.2 (Satır 2): PANO_adi1, en_buyuk_top_kesit, Sonuç=UYGUN
        - 6.3 (Satır 6): PANO_adi1, en, boy, izo_direnci, izo_uygunluk
        """
        if not self.doc or len(self.doc.tables) < 4:
            return

        table = self.doc.tables[3]  # Tablo 3 (index 3)

        # === PANO ADI ===
        pano_adi = data.get('pano_adi', '')
        print(f"[DEBUG] fill_potansiyel_dengeleme: pano_adi = '{pano_adi}'")

        # === EN BÜYÜK TOPRAKLAMA KESİTİ ===
        en_buyuk_top_kesit = ""
        if fonksiyon_testleri:
            max_kesit = 0
            for row in fonksiyon_testleri:
                kesit_str = row.get('Toprak Kesiti', '') or row.get('PE Kesiti', '')
                if kesit_str:
                    try:
                        kesit = str(kesit_str).replace(',', '.').lower()
                        if 'x' in kesit:
                            parts = kesit.split('x')
                            kesit_val = float(parts[-1])
                        else:
                            kesit_val = float(kesit)
                        if kesit_val > max_kesit:
                            max_kesit = kesit_val
                            en_buyuk_top_kesit = str(kesit_str)
                    except:
                        pass

        print(f"[INFO] 6.2 Potansiyel Dengeleme: pano_adi={pano_adi}, en_buyuk_top_kesit={en_buyuk_top_kesit}")

        # === 6.3 ZEMİN İZOLASYONU DEĞERLERİ ===
        gozle_kontrol = data.get('gozle_kontrol', {})
        zemin_izolasyonu = gozle_kontrol.get('Zemin Izolasyonu') or \
                           gozle_kontrol.get('kontroller', {}).get('Zemin Izolasyonu', 'Uygun')

        print(f"[INFO] 6.3 Zemin İzolasyonu durumu: {zemin_izolasyonu}")

        if zemin_izolasyonu == 'Uygun':
            en_val, boy_val, izo_direnci, izo_uygunluk = "1", "1", ">50MΩ", "UYGUN"
        elif zemin_izolasyonu in ['Uygun Degil', 'Uygun Değil']:
            en_val, boy_val, izo_direnci, izo_uygunluk = "x", "x", "-", "UYGUN DEĞİL"
        else:  # Uygulanamaz
            en_val, boy_val, izo_direnci, izo_uygunluk = "-", "-", "-", "UYGULANAMAZ"

        # === TÜM PLACEHOLDER'LARI DEĞİŞTİR ===
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                cell_text_lower = cell_text.lower()

                # PANO_adi1 (6.2 satır 2 ve 6.3 satır 6)
                if 'pano_adi1' in cell_text_lower:
                    self._set_cell_value(cell, pano_adi, font_size=9)

                # en_buyuk_top_kesit (6.2)
                if 'en_buyuk_top_kesit' in cell_text_lower:
                    self._set_cell_value(cell, en_buyuk_top_kesit or "6", font_size=9)

                # en (6.3 - tek kelime)
                if cell_text.strip() == 'en':
                    self._set_cell_value(cell, en_val, font_size=9)

                # boy (6.3 - tek kelime)
                if cell_text.strip() == 'boy':
                    self._set_cell_value(cell, boy_val, font_size=9)

                # izo_direnci veya İzo_direnci (6.3) - Türkçe İ için özel kontrol
                if 'izo_direnci' in cell_text or 'İzo_direnci' in cell_text:
                    self._set_cell_value(cell, izo_direnci, font_size=9)

                # izo_uygunluk veya İzo_uygunluk (6.3) - Türkçe İ için özel kontrol
                if 'izo_uygunluk' in cell_text or 'İzo_uygunluk' in cell_text:
                    self._set_cell_value(cell, izo_uygunluk, font_size=9)

        # === DOĞRUDAN HÜCRE İNDEKSLERİ İLE DEĞİŞTİRME (6.3 Satır 6) ===
        # Şablon: R6C4=izo_direnci, R6C5-6=izo_uygunluk (merged)
        if len(table.rows) > 6:
            row = table.rows[6]
            # izo_direnci (R6C4)
            if len(row.cells) > 4:
                cell4 = row.cells[4]
                self._set_cell_value(cell4, izo_direnci, font_size=9)
            # izo_uygunluk (R6C5, R6C6 merged olabilir)
            if len(row.cells) > 5:
                cell5 = row.cells[5]
                self._set_cell_value(cell5, izo_uygunluk, font_size=9)
            if len(row.cells) > 6:
                cell6 = row.cells[6]
                if cell6 != row.cells[5]:  # Farklı hücre ise
                    self._set_cell_value(cell6, izo_uygunluk, font_size=9)

    def fill_ana_dagitim_pano(self, data: Dict[str, Any], pano_adi: str = None):
        """Ana Dağıtım Pano bilgilerini doldurur (Tablo 1 - 2.1 DETAY BİLGİLER).

        Satır indeksleri (0-indexed):
        - Satır 10: Enerji sağlayan kuruluş (sütun 1), Şebeke tipi (sütun 7 - checkbox)
        - Satır 13: Ekipman kullanım amacı (sütun 7), Son kontrol tarihi (sütun 12)
        - Satır 14-17: Faz iletkenleri, Topraklama direnci, İletken kesitleri
        - Satır 18: RCD bilgisi (sütun 12)
        - Satır 19: RCD test bilgisi (sütun 12)
        """
        if not self.doc or not self.doc.tables:
            return

        table = self.doc.tables[0]

        # Pano adı (Ekipman Kullanım Amacı - Satır 13, Sütun 7)
        if pano_adi:
            cell = self._get_cell_safe(table, 13, 7)
            self._set_cell_value(cell, pano_adi)

        # Metin alanları doldurulur
        field_mapping = {
            # Satır 10 - Enerji sağlayan kuruluş
            'Enerji Saglayan Kurulus': (10, 1),

            # Satır 14 - Temel topraklama direnci (sütun 9'da [BOS] alanı)
            'Temel Topraklama Direnci (Ohm)': (14, 9),

            # Satır 15 - İlave topraklama elektrotu
            'Ilave Topraklama Elektrotu Detaylari': (15, 9),

            # Satır 16 - Sistem topraklama iletkeni kesiti
            'Sistem Topraklama Iletkeni Kesiti (mm2)': (16, 9),

            # Satır 17 - Ana eşpotansiyel iletkeni kesiti
            'Ana Espotansiyel Iletkeni Kesiti (mm2)': (17, 9),

            # Satır 18 - RCD bilgisi (placeholder tabanlı)
            'Ana RCD Tipi': (18, 12),

            # Satır 19 - RCD test bilgisi (sütun 12)
            'Ana RCD Test Akimi (mA)': (19, 12),
        }

        for field_name, (row, col) in field_mapping.items():
            if field_name in data:
                cell = self._get_cell_safe(table, row, col)
                value = data[field_name]

                # RCD alanları için birleştirme
                if field_name == 'Ana RCD Tipi':
                    rcd_tipi = data.get('Ana RCD Tipi', '')
                    rcd_anma = data.get('Ana RCD Anma Akimi (A)', '')
                    if rcd_tipi and rcd_anma:
                        value = f"{rcd_tipi} {rcd_anma} A"
                    elif rcd_anma:
                        value = f"{rcd_anma} A"

                elif field_name == 'Ana RCD Test Akimi (mA)':
                    test_akim = data.get('Ana RCD Test Akimi (mA)', '')
                    acma_suresi = data.get('Ana RCD Acma Suresi (ms)', '')
                    if test_akim and acma_suresi:
                        value = f"{test_akim} mA / {acma_suresi} ms"
                    elif test_akim:
                        value = f"{test_akim} mA"

                self._set_cell_value(cell, value)

        # === PLACEHOLDER TABANLI DEĞİŞTİRME ===
        # Tablo 1'deki tüm placeholder'ları Excel verisiyle değiştir

        # Z_E'den I_f otomatik hesapla: I_f = 230 / Z_E
        z_e_value = data.get('Dis Cevrim Empedansi Z_E (Ohm)', '')
        i_f_value = ''
        if z_e_value:
            try:
                z_e_float = float(str(z_e_value).replace(',', '.'))
                if z_e_float > 0:
                    i_f_calculated = 230 / z_e_float  # Ampere
                    if i_f_calculated >= 1000:
                        i_f_value = f"{i_f_calculated/1000:.2f}"  # kA olarak
                    else:
                        i_f_value = f"{i_f_calculated:.1f}"  # A olarak
            except (ValueError, ZeroDivisionError):
                pass

        # Şebeke tipi için checkbox işaretleme
        sebeke_tipi = data.get('Sebeke Tipi', '')
        print(f"[DEBUG] Şebeke Tipi değeri: '{sebeke_tipi}'")

        # Şebeke tipi placeholder mapping (placeholder -> tip adı)
        # ÖNEMLİ: Uzun placeholder'lar önce gelmeli (t_n_c-s önce, t_n sonra)
        # Şablonda: t_t TT, i_t IT, t_n TN, t_n_c-s TN-CS, t_n-c TN-C, t_n-s TN-S
        sebeke_placeholders = [
            ('t_n_c-s', 'TN-CS'),  # Önce uzun olanlar
            ('t_n-c', 'TN-C'),
            ('t_n-s', 'TN-S'),
            ('t_t', 'TT'),          # Sonra kısa olanlar
            ('i_t', 'IT'),
            ('t_n', 'TN'),          # t_n en sonda (diğerlerini bozmasın)
        ]

        placeholder_mapping = {
            'ADP_tip': data.get('Ana Kesici Tipi', ''),  # Ana Kesici Tipi
            'ADP_anma': data.get('Ana Kesici Nominal Akimi', ''),  # Ana Kesici Nominal Akım
            'RCD_tipi': data.get('Ana RCD Tipi', ''),  # RCD Tipi
            'RCD_anma': data.get('Ana RCD Anma Akimi (A)', ''),  # RCD Anma Akımı (A)
            'Z_E': str(z_e_value) if z_e_value else '',  # Dış çevrim empedansı
            'I_f': i_f_value,  # Hata akımı (otomatik hesaplanan)
            'Sistem_top': data.get('Sistem Topraklama Kesiti (mm2)', ''),  # Sistem topraklama kesiti
            'Ana_top': data.get('Ana Espotansiyel Kesiti (mm2)', ''),  # Ana eşpotansiyel kesiti
        }

        # Tablo 1'deki tüm hücrelerde placeholder'ları değiştir
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                for placeholder, value in placeholder_mapping.items():
                    if value and placeholder in cell_text:
                        print(f"[DEBUG] Tablo 1'de bulundu: '{placeholder}' in cell '{cell_text[:50]}'")
                        self._replace_placeholder_in_cell(cell, placeholder, str(value))

        # Şebeke tipi placeholder'larını TÜM tablolarda ara (sıralı liste - uzun olanlar önce)
        print(f"[DEBUG] Toplam tablo sayısı: {len(self.doc.tables)}")
        for tbl_idx, tbl in enumerate(self.doc.tables):
            for row in tbl.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    # Sıralı listeyi kullan (uzun placeholder'lar önce)
                    for placeholder, tip in sebeke_placeholders:
                        if placeholder in cell_text:
                            value = '[X]' if sebeke_tipi == tip else '[ ]'
                            print(f"[DEBUG] Tablo {tbl_idx}'da bulundu: '{placeholder}' -> '{value}'")
                            self._replace_placeholder_in_cell(cell, placeholder, value)

        # === SPD (Aşırı Gerilim Koruma Cihazı) checkbox'ı ===
        # Parafudr Imax değeri doluysa Evet, boşsa veya "-" ise Hayır
        parafudr_imax = str(data.get('Parafudr Imax (kA)', '')).strip()
        parafudr_var = parafudr_imax and parafudr_imax not in ('', '-', 'x', 'X', 'yok', 'Yok', 'YOK')
        spd_evet = '[X]' if parafudr_var else '[ ]'
        spd_hayir = '[ ]' if parafudr_var else '[X]'

        # Tüm tablolarda SPD placeholder'larını ara
        for tbl in self.doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if 'spd_evet' in cell.text:
                        self._replace_placeholder_in_cell(cell, 'spd_evet', spd_evet)
                    if 'spd_hayir' in cell.text:
                        self._replace_placeholder_in_cell(cell, 'spd_hayir', spd_hayir)

    def fill_cihaz_bilgileri(self, data: Dict[str, Any]):
        """Cihaz bilgilerini doldurur (Tablo 1 - 3. ve 4. bölümler)."""
        if not self.doc or not self.doc.tables:
            return

        table = self.doc.tables[0]

        # Termal Kamera
        if 'termal_kamera' in data:
            tk = data['termal_kamera']
            mapping = {
                'Cihaz Adi': (26, 3),
                'Kalibrasyon Tarihi': (27, 3),
                'Gecerlilik Tarihi': (28, 3),
                'Seri No': (29, 3),
                'Kalibrasyon No': (30, 3),
            }
            for field, (row, col) in mapping.items():
                if field in tk:
                    cell = self._get_cell_safe(table, row, col)
                    self._set_cell_value(cell, tk[field])

        # Ölçüm Aleti
        if 'olcum_aleti' in data:
            oa = data['olcum_aleti']
            mapping = {
                'Cihaz Adi': (32, 3),
                'Kalibrasyon Tarihi': (33, 3),
                'Gecerlilik Tarihi': (34, 3),
                'Seri No': (35, 3),
                'Kalibrasyon No': (36, 3),
            }
            for field, (row, col) in mapping.items():
                if field in oa:
                    cell = self._get_cell_safe(table, row, col)
                    self._set_cell_value(cell, oa[field])

    def fill_gozle_kontrol(self, data: Dict[str, Any]):
        """Gözle kontrol bölümünü doldurur (Tablo 2 - TÜM MADDELER)."""
        if not self.doc or len(self.doc.tables) < 2:
            return

        table = self.doc.tables[1]  # 2. tablo

        # Pano adı
        if 'pano_adi' in data:
            cell = self._get_cell_safe(table, 2, 1)
            self._set_cell_value(cell, data['pano_adi'])

        # Kontroller - Excel'deki her satırı rapordaki karşılığına eşle
        if 'kontroller' in data:
            kontroller = data['kontroller']

            # Detaylı eşleme - Excel alan adı: (tablo satırı, sütun)
            kontrol_mapping = {
                'Kablo Sebeke Tarafi': (5, 1),
                'Kablo Donanim Tarafi': (5, 3),
                'Pano Sabitlenmesi (Depreme Dayaniklilik)': (6, 1),
                'Dis Darbelere Karsi Koruma Onlemi': (6, 3),
                'Elektrik Panosu Etrafinda Yabanci Malzemeler': (7, 1),
                'Zemin Izolasyonu': (7, 3),
                'Topraklama Iletkeni': (9, 1),
                'Ana Potansiyel Dengeleme Iletkeni': (9, 3),
                'Ek Potansiyel Dengeleme Iletkeni (Tamamlayici)': (10, 1),
                'Ek Potansiyel Dengeleme Iletkeni': (10, 1),  # Alternatif isim
                'Pano Kapak Baglantisi Kontrolu 6 mm2': (10, 3),
                'Elektriksel Olmayan Tesislere Yaklasma ve Diger Etkilerin Kontrolu': (12, 1),
                'Bant I ve Bant II Ayrilmasi, Bant II Yalitimi': (12, 3),
                'Bant Ayrilmasi': (12, 3),  # Kısa isim
                'Guvenlik Devre Ayrilmasi': (13, 1),
                'Pano Ic Kapak, Faza Erisim Engeli veya Pleksi Koruma': (13, 3),
                'Pano Ic Kapak': (13, 3),  # Kısa isim
                'Semalar, Talimatlar, Devre Cizimleri ve Kisa Bilgiler': (15, 1),
                'Koruma Cihaz ve Terminal Etiket': (15, 3),
                'Tehlike Isaretleri ve Diger Uyari Isaretleri': (16, 1),
                'Tehlike Isaretleri': (16, 1),  # Kısa isim
                'Kablo Yollarinin Uygunlugu ve Mekanik Koruma': (18, 1),
                'Kablo Yollari': (18, 1),  # Kısa isim
                'Kablo Renk Kodlari (Notr: Mavi, Toprak: Sari/Yesil)': (18, 3),
                'Kablo Renk Kodlari': (18, 3),  # Kısa isim
                'Tesisat Yontemi': (19, 1),
                'Yangin Engeli, Uygun Kilitleme ve Sicaklik Etkisine Karsi Koruma': (19, 3),
                'Yangin Engeli': (19, 3),  # Kısa isim
                # NOT: Fotograf Tarihi/No ve Acil Durum Aydinlatma burada yok!
                # Bunlar Fluke'dan gelecek ve GK_24/25/29 placeholder ile değiştirilecek
                # _set_cell_value kullanmak placeholder'ı siler, GK_XX değiştirme çalışmaz
                'Kontak Gevsekligi Isinmasi': (21, 3),
                'Asiri Yuk Isinmasi (PVC Kablolar Icin >70 derece)': (22, 3),
                'Asiri Yuk Isinmasi': (22, 3),  # Kısa isim
                'Ekipman Yakininda Elektriksel Ekipman Yangin Sondurme Tertibati': (24, 1),
                'Yangin Sondurme': (24, 1),  # Kısa isim
                'Ekipman Temizlik/Bakim Durumu': (24, 3),
                'Ekipman Temizlik': (24, 3),  # Kısa isim
                'Pano Ici ve Baglantilarinin Korozyon Kontrolu': (25, 1),
                'Korozyon Kontrolu': (25, 1),  # Kısa isim
                # NOT: Acil Durum Aydinlatma da burada yok - GK_29 olarak değiştirilecek
            }

            for field, (row, col) in kontrol_mapping.items():
                if field in kontroller:
                    cell = self._get_cell_safe(table, row, col)
                    self._set_cell_value(cell, kontroller[field])

        # === GK_XX PLACEHOLDER DEĞİŞTİRME ===
        # Gözle Kontrol tablosundaki GK_01 - GK_29 placeholder'larını değiştir
        if 'kontroller' in data:
            kontroller = data['kontroller']

            # GUI alan adı -> GK_XX placeholder eşleme (TAM 29 ALAN)
            gk_placeholder_mapping = {
                'Kablo Sebeke Tarafi': 'GK_01',
                'Pano Sabitlenmesi': 'GK_02',
                'Elektrik Panosu Etrafinda Yabanci Malzemeler': 'GK_03',
                'Kablo Donanim Tarafi': 'GK_04',
                'Dis Darbelere Karsi Koruma Onlemi': 'GK_05',
                'Zemin Izolasyonu': 'GK_06',
                'Topraklama Iletkeni': 'GK_07',
                'Ek Potansiyel Dengeleme Iletkeni': 'GK_08',
                'Ana Potansiyel Dengeleme Iletkeni': 'GK_09',
                'Pano Kapak Baglantisi Kontrolu 6 mm2': 'GK_10',
                'Elektriksel Olmayan Tesislere Yaklasma': 'GK_11',
                'Guvenlik Devre Ayrilmasi': 'GK_12',
                'Bant Ayrilmasi': 'GK_13',
                'Pano Ic Kapak': 'GK_14',
                'Semalar Talimatlar': 'GK_15',
                'Tehlike Isaretleri': 'GK_16',
                'Koruma Cihaz ve Terminal Etiket': 'GK_17',
                'Kablo Yollari': 'GK_18',
                'Tesisat Yontemi': 'GK_19',
                'Kablo Renk Kodlari': 'GK_20',
                'Yangin Engeli': 'GK_21',
                'Kontak Gevsekligi Isinmasi': 'GK_22',
                'Asiri Yuk Isinmasi': 'GK_23',
                'Fotograf Tarihi': 'GK_24',
                'Fotograf No': 'GK_25',
                'Yangin Sondurme': 'GK_26',
                'Korozyon Kontrolu': 'GK_27',
                'Ekipman Temizlik': 'GK_28',
                'Acil Durum Aydinlatma': 'GK_29',
            }

            # Tablo 2'deki tüm hücrelerde GK_XX placeholder'ları değiştir
            for row in table.rows:
                for cell in row.cells:
                    for field_name, placeholder in gk_placeholder_mapping.items():
                        if field_name in kontroller:
                            value = kontroller[field_name]
                            if value:
                                self._replace_placeholder_in_cell(cell, placeholder, str(value))

    def fill_fonksiyon_testleri(self, data: List[Dict[str, Any]], ana_pano_data: Dict[str, Any] = None):
        """Fonksiyon testleri bölümünü doldurur (Tablo 3).

        ana_pano_data: AnaDagitimPano verisi - placeholder değerleri için kullanılır
        """
        if not self.doc or len(self.doc.tables) < 3:
            return

        table = self.doc.tables[2]  # 3. tablo

        # === PLACEHOLDER DEĞİŞTİRME (Tablo 3) ===
        if ana_pano_data:
            # Z_ln'den 380/Z_ln otomatik hesapla
            z_ln_value = ana_pano_data.get('Faz-Notr Cevrim Empedansi Z_ln (Ohm)', '')
            kisa_devre = ''
            if z_ln_value:
                try:
                    z_ln_float = float(str(z_ln_value).replace(',', '.'))
                    if z_ln_float > 0:
                        kd = 380 / z_ln_float  # Amper
                        kisa_devre = f"{int(round(kd))}"
                except (ValueError, ZeroDivisionError):
                    pass

            # Önce 380/Z_ln, sonra diğerleri (Z_ln önce yazılırsa 380/Z_ln bozuluyor)
            placeholder_items = [
                ('380/Z_ln', kisa_devre),
                ('PANO_adi1', ana_pano_data.get('Pano Adi (PANO_adi1)', '')),
                ('Z_x', ana_pano_data.get('Faz-Toprak Cevrim Empedansi Z_x (Ohm)', '')),
                ('Z_ln', z_ln_value),
                ('F_F', ana_pano_data.get('Gerilim F-F (V)', '')),
                ('L_N', ana_pano_data.get('Gerilim L-N (V)', '')),
                ('N_PE', ana_pano_data.get('Gerilim N-PE (V)', '')),
                ('PARAFUDR_TIP', ana_pano_data.get('Parafudr Tipi', '')),
                ('PARAFUDR_Imax', ana_pano_data.get('Parafudr Imax (kA)', '')),
            ]

            # Tablo 3'teki tüm hücrelerde placeholder'ları değiştir
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholder_items:
                        if value:
                            self._replace_placeholder_in_cell(cell, placeholder, str(value))

        # Veri satırları 10'dan başlıyor (0-indexed)
        start_row = 10
        existing_rows = len(table.rows)

        # Sütun eşlemesi - DOCX template yapısına göre (0-indexed)
        # Sütun 0: No.
        # Sütun 1-2: Linye Adı (merged)
        # Sütun 3: Açma eğrisi
        # Sütun 4: Kutup sayısı
        # Sütun 5-6: In(A) (merged)
        # Sütun 7: Icu
        # Sütun 8-9: Faz kesiti (merged)
        # Sütun 10: N/PEN Kesiti
        # Sütun 11-12: PE Kesiti (merged)
        # Sütun 13-14: Ib (merged)
        # Sütun 15: Iz
        # Sütun 16-17: IΔ mA (merged)
        # Sütun 18: TΔ ms
        # Sütun 19: Sonuç
        col_mapping = {
            'Linye Adi': 1,       # Merged hücre 1-2
            'Acma Egrisi': 3,
            'Kutup Sayisi': 4,
            'In (A)': 5,          # Merged hücre 5-6
            'Icu': 7,
            'Faz Kesiti': 8,      # Merged hücre 8-9
            'Notr Kesiti': 10,
            'Toprak Kesiti': 11,  # Merged hücre 11-12
            'Ib': 13,             # Merged hücre 13-14
            'Iz': 15,
            'RCD mA': 16,         # Merged hücre 16-17
            'RCD ms': 18,
            'Sonuc': 19,
        }

        # Kablo Akım Taşıma Kapasiteleri (Grup 2)
        iz_table = const.IZ_TABLE

        def parse_kesit(section_val: str) -> float:
            if section_val is None:
                return 0.0
            val = str(section_val).lower().replace(',', '.').strip()
            if not val:
                return 0.0
            factor = 1
            base = val
            if 'x' in val:
                try:
                    f_str, base = val.split('x', 1)
                    factor = float(f_str)
                except ValueError:
                    factor = 1
            try:
                size = float(base)
            except ValueError:
                return 0.0
            return size * factor

        def calc_iz(section_val: str) -> str:
            size = parse_kesit(section_val)
            if not size:
                return ""
            if size in iz_table:
                return f"{iz_table[size]:.0f}"
            for k in sorted(iz_table.keys()):
                if size <= k:
                    return f"{iz_table[k]:.0f}"
            return f"{list(iz_table.values())[-1]:.0f}"

        # Şablondaki son veri satırının indeksini sakla (format kopyalamak için)
        template_last_row = start_row  # Şablonda en az 1 veri satırı olduğunu varsayıyoruz
        if len(table.rows) > start_row:
            template_last_row = min(start_row + 21, len(table.rows) - 1)  # 22 satırlık şablon

        for i, row_data in enumerate(data):
            row_idx = start_row + i

            # Iz Otomatik Hesapla (Faz Kesiti veya gelen Iz (A))
            if not row_data.get('Iz'):
                if 'Iz (A)' in row_data and row_data.get('Iz (A)'):
                    row_data['Iz'] = row_data.get('Iz (A)')
                else:
                    row_data['Iz'] = calc_iz(row_data.get('Faz Kesiti'))

            # Eğer yeterli satır yoksa, şablon satırını kopyalayarak yeni satır ekle
            while row_idx >= len(table.rows):
                # Şablondaki son veri satırını kopyala (format korunur)
                self._copy_row(table, template_last_row)

            # Satır numarasını güncelle
            cell = self._get_cell_safe(table, row_idx, 0)
            if cell:
                self._set_cell_value(cell, str(i + 1), font_size=6)

            # Icu - varsa kullan, yoksa 6kA default
            icu_cell = self._get_cell_safe(table, row_idx, 7)
            icu_val = row_data.get('Icu', '') or "6"
            self._set_cell_value(icu_cell, str(icu_val), font_size=6)

            for field, col in col_mapping.items():
                val = row_data.get(field, '')

                # Sonuç boşsa varsayılan olarak 'Uygun' yazsın (Görünmeme ihtimaline karşı)
                if field == 'Sonuc' and not val:
                    val = "Uygun"

                if val is not None:
                    # '-' was previously treated empty; now allow markers like 'x'
                    val_to_write = "" if val in (None, "") else val
                    cell = self._get_cell_safe(table, row_idx, col)

                    # KRİTİK DÜZELTME: Sonuç sütunu (21) standart yöntemle boş kalabiliyor.
                    if col == 21 and cell:
                        cell.text = str(val_to_write)
                        # Fontu manuel küçült
                        try:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(6)
                        except:
                            pass
                    else:
                        self._set_cell_value(cell, str(val_to_write), font_size=6)

    def add_thermal_images(self, image_paths: List[str]):
        """Ekipman Fotoğrafları bölümüne termal görüntüleri ekler (Tablo 5 - index 5)."""
        if not self.doc or len(self.doc.tables) < 6:
            return

        table = self.doc.tables[5]  # Tablo 5 (index 5) = 8. EKİPMAN FOTOĞRAFLARI

        # Fotoğraf hücresi (satır 1)
        if len(table.rows) > 1 and image_paths:
            cell = table.rows[1].cells[0]

            # Hücreyi temizle
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""

            # Görüntüleri ekle (son 2 görsel)
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

            for img_path in image_paths:  # Zaten son 2 görsel gelecek
                if os.path.exists(img_path):
                    run = para.add_run()
                    run.add_picture(img_path, width=Inches(2.5))  # ~%50 boyut
                    run.add_text("  ")  # Görüntüler arası boşluk

    def generate_kusur_aciklamasi(self, kusurlar: List[Dict[str, str]]) -> str:
        """Kusur listesinden açıklama metni oluşturur."""
        if not kusurlar:
            return ""

        kusur_texts = []
        for kusur in kusurlar:
            derece = kusur.get('derece', '*')
            madde = kusur.get('madde', '')
            kusur_texts.append(f"{derece} {madde}")

        return "\n".join(kusur_texts)

    def fill_sonuc(self, data: Dict[str, Any], kusurlar: List[Dict[str, str]] = None):
        """Sonuç ve kanaat bölümünü doldurur."""
        if not self.doc or len(self.doc.tables) < 5:
            return

        # Tablo 4 (index 4) - 7. KUSUR AÇIKLAMALARI
        if len(self.doc.tables) >= 5:
            table = self.doc.tables[4]
            if len(table.rows) > 1:
                cell = self._get_cell_safe(table, 1, 0)

                # Kusur açıklamasını oluştur
                if kusurlar:
                    kusur_metni = self.generate_kusur_aciklamasi(kusurlar)
                    self._set_cell_value(cell, kusur_metni)
                else:
                    self._set_cell_value(cell, "Herhangi bir kusur tespit edilmemiştir.")

        # Tablo 5 - Ek Notlar
        if 'Ek Notlar' in data and len(self.doc.tables) >= 5:
            # Notlar bölümü için tablo 5'in ilgili hücresine yaz
            pass

    def fill_uygunluk(self, uygunluk_durumu: str):
        """10. SONUÇ VE KANAAT bölümündeki 'uygunluk' placeholder'ını doldurur.

        Args:
            uygunluk_durumu: "Uygun" veya "Uygun Değil"
        """
        if not self.doc:
            return

        # Uygunluk metnini belirle
        if uygunluk_durumu == "Uygun":
            uygunluk_metni = "kullanımı uygundur"
        else:
            uygunluk_metni = "kullanımı uygun değildir"

        print(f"[INFO] Uygunluk placeholder değiştiriliyor: '{uygunluk_durumu}' -> '{uygunluk_metni}'")

        def replace_in_paragraph(para):
            """Paragraftaki 'uygunluk' placeholder'ını değiştirir, sadece yeni metni formatlar."""
            full_text = para.text
            if 'uygunluk' not in full_text:
                return False

            # Placeholder'ın konumunu bul
            idx = full_text.find('uygunluk')
            before_text = full_text[:idx]
            after_text = full_text[idx + len('uygunluk'):]

            # Mevcut paragrafın ilk run'ının formatını sakla (varsayılan format olarak)
            original_font_name = None
            original_font_size = None
            original_bold = None
            if para.runs:
                first_run = para.runs[0]
                original_font_name = first_run.font.name
                original_font_size = first_run.font.size
                original_bold = first_run.font.bold

            # Tüm run'ları temizle
            for run in list(para.runs):
                run._element.getparent().remove(run._element)

            # Yeni run'ları ekle
            # 1. Önceki metin (orijinal format)
            if before_text:
                run_before = para.add_run(before_text)
                if original_font_name:
                    run_before.font.name = original_font_name
                if original_font_size:
                    run_before.font.size = original_font_size
                if original_bold is not None:
                    run_before.font.bold = original_bold

            # 2. Uygunluk metni (7.5pt bold)
            run_uygunluk = para.add_run(uygunluk_metni)
            run_uygunluk.font.name = 'Calibri'
            run_uygunluk.font.size = Pt(7.5)
            run_uygunluk.font.bold = True

            # 3. Sonraki metin (orijinal format)
            if after_text:
                run_after = para.add_run(after_text)
                if original_font_name:
                    run_after.font.name = original_font_name
                if original_font_size:
                    run_after.font.size = original_font_size
                if original_bold is not None:
                    run_after.font.bold = original_bold

            return True

        # Tüm tablolarda ara
        for tbl_idx, table in enumerate(self.doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if replace_in_paragraph(para):
                            print(f"[DEBUG] Tablo {tbl_idx}'de 'uygunluk' değiştirildi")

        # Paragraflar içinde de ara (tablo dışındaki alanlar için)
        for para in self.doc.paragraphs:
            if replace_in_paragraph(para):
                print(f"[DEBUG] Paragrafta 'uygunluk' değiştirildi")


    def add_page_break_before_table(self, table_index: int = 3):
        """Belirtilen tablodan önce sayfa sonu ekler.

        Args:
            table_index: Sayfa sonu eklenecek tablonun indeksi (varsayılan: 3 = 6.2 Potansiyel Dengeleme)
        """
        if not self.doc or len(self.doc.tables) <= table_index:
            return

        table = self.doc.tables[table_index]

        # Tablonun XML elementini al
        tbl = table._tbl

        # Tablo öncesine sayfa sonu içeren paragraf ekle
        # Yeni paragraf oluştur
        p = OxmlElement('w:p')

        # Paragraf properties
        pPr = OxmlElement('w:pPr')
        p.append(pPr)

        # Run oluştur
        r = OxmlElement('w:r')

        # Sayfa sonu (page break)
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)

        p.append(r)

        # Paragrafı tablonun önüne ekle
        tbl.addprevious(p)

        print(f"[INFO] Tablo {table_index} öncesine sayfa sonu eklendi")

    def save(self, output_path: str):
        """Belgeyi kaydeder."""
        if not self.doc:
            raise ValueError("Belge yuklu degil")

        # Çıktı klasörünü oluştur
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        self.doc.save(output_path)
        return output_path


# Test için
if __name__ == "__main__":
    template_path = "Elektrik Tesisatı Gözle Kontrol ve Fonksiyon Testleri Periyodik Kontrol Raporu.docx"

    if os.path.exists(template_path):
        writer = DocxWriter(template_path)
        writer.load_template()

        # Test verisi
        test_data = {
            'Firma Adi': 'TEST FIRMA A.S.',
            'Rapor Numarasi': '2025-001',
            'Rapor Tarihi': '12.12.2025',
        }

        writer.fill_firma_bilgileri(test_data)

        output_path = "test_output.docx"
        writer.save(output_path)
        print(f"Test raporu olusturuldu: {output_path}")
    else:
        print(f"Test sablonu bulunamadi: {template_path}")
