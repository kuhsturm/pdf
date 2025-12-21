"""
Fluke DOCX Termal Görüntü Çıkarıcı
Fluke termal kamera raporlarından görüntü ve verileri çıkarır.
"""

from docx import Document
from docx.shared import Inches
import os
import shutil
import zipfile
import tempfile


class FlukeExtractor:
    """Fluke DOCX dosyasından termal görüntüleri ve verileri çıkarır."""

    def __init__(self, fluke_docx_path: str):
        """
        Args:
            fluke_docx_path: Fluke DOCX dosyasının yolu
        """
        self.docx_path = fluke_docx_path
        self.doc = None
        self.images = []
        self.temperature_data = {}
        self.device_info = {}

    def load(self) -> bool:
        """DOCX dosyasını yükler."""
        if not os.path.exists(self.docx_path):
            raise FileNotFoundError(f"Dosya bulunamadi: {self.docx_path}")

        self.doc = Document(self.docx_path)
        return True

    def extract_images(self, output_dir: str, only_last_two: bool = True) -> list:
        """
        DOCX içindeki görüntüleri çıkarır.

        Args:
            output_dir: Görüntülerin kaydedileceği klasör
            only_last_two: Sadece son 2 görüntüyü al (termal için visible + infrared)

        Returns:
            Çıkarılan görüntü dosya yollarının listesi
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        extracted_images = []

        # DOCX aslında bir ZIP dosyasıdır
        with zipfile.ZipFile(self.docx_path, 'r') as zip_ref:
            for file_info in zip_ref.filelist:
                if file_info.filename.startswith('word/media/') and \
                   file_info.filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                    # Dosyayı çıkar
                    image_data = zip_ref.read(file_info.filename)

                    # Dosya adını al
                    image_name = os.path.basename(file_info.filename)
                    output_path = os.path.join(output_dir, image_name)

                    with open(output_path, 'wb') as f:
                        f.write(image_data)

                    extracted_images.append(output_path)

        # Küçük/logolu görselleri ele (örn. FLUKE logo). 30KB altını at.
        filtered = [p for p in extracted_images if os.path.getsize(p) >= 30_000]
        extracted_images = filtered if filtered else extracted_images

        # Sadece son 2 görseli al (visible ve infrared termal görüntüler)
        if only_last_two and len(extracted_images) > 2:
            extracted_images = extracted_images[-2:]

        self.images = extracted_images
        return extracted_images

    def extract_temperature_data(self) -> dict:
        """
        Sıcaklık verilerini çıkarır.

        Returns:
            Sıcaklık verileri sözlüğü
        """
        if not self.doc:
            self.load()

        temp_data = {
            'merkez_nokta': None,
            'maksimum': None,
            'minimum': None,
            'ortam_sicakligi': None
        }

        # Tablo 3'ten sıcaklık verilerini oku
        for table in self.doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                # Merkez Nokta, Maksimum, Min. değerlerini ara
                if 'Merkez Nokta' in cells:
                    try:
                        idx = cells.index('Merkez Nokta')
                        if idx + 1 < len(cells):
                            temp_data['merkez_nokta'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                if 'Maksimum' in cells:
                    try:
                        idx = cells.index('Maksimum')
                        if idx + 1 < len(cells):
                            temp_data['maksimum'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                if 'Min.' in cells:
                    try:
                        idx = cells.index('Min.')
                        if idx + 1 < len(cells):
                            temp_data['minimum'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                # Ortam sıcaklığı
                if 'Ortam' in cells:
                    try:
                        idx = cells.index('Ortam')
                        if idx + 1 < len(cells):
                            temp_data['ortam_sicakligi'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

        self.temperature_data = temp_data
        return temp_data

    def extract_device_info(self) -> dict:
        """
        Cihaz bilgilerini çıkarır.

        Returns:
            Cihaz bilgileri sözlüğü
        """
        if not self.doc:
            self.load()

        device_info = {
            'rapor_no': None,
            'cihaz_modeli': None,
            'seri_no': None,
            'emisivite': None,
            'mesafe': None,
            'tarih_saat': None,
            'foto_no': None,  # Fotoğraf no. (GK_25 gibi)
            'foto_tarihi': None,  # Fotoğraf tarihi (2025-11-11 gibi)
        }

        for table in self.doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                # Rapor No
                if 'Rapor No' in cells:
                    try:
                        idx = cells.index('Rapor No')
                        if idx + 1 < len(cells):
                            device_info['rapor_no'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                # Ekipman bilgileri (TC01A, SN, vb.)
                if 'Ekipman' in cells:
                    try:
                        idx = cells.index('Ekipman')
                        if idx + 1 < len(cells):
                            device_info['cihaz_modeli'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                if 'SN' in cells:
                    try:
                        idx = cells.index('SN')
                        if idx + 1 < len(cells):
                            device_info['seri_no'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                if 'Emisivite' in cells:
                    try:
                        idx = cells.index('Emisivite')
                        if idx + 1 < len(cells):
                            device_info['emisivite'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                if 'Mesafe' in cells:
                    try:
                        idx = cells.index('Mesafe')
                        if idx + 1 < len(cells):
                            device_info['mesafe'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                if 'Saat' in cells:
                    try:
                        idx = cells.index('Saat')
                        if idx + 1 < len(cells):
                            device_info['tarih_saat'] = cells[idx + 1]
                    except (ValueError, IndexError):
                        pass

                # Fotoğraf no. ("Fotoğraf no." veya "Photo no." vb.)
                for foto_label in ['Fotoğraf no.', 'Fotoğraf no', 'Photo no.', 'Photo no', 'Foto no.', 'Foto no']:
                    if foto_label in cells:
                        try:
                            idx = cells.index(foto_label)
                            if idx + 1 < len(cells):
                                device_info['foto_no'] = cells[idx + 1]
                        except (ValueError, IndexError):
                            pass

                # Fotoğraf tarihi
                for tarih_label in ['Fotoğraf tarihi', 'Photo date', 'Tarih']:
                    if tarih_label in cells:
                        try:
                            idx = cells.index(tarih_label)
                            if idx + 1 < len(cells):
                                device_info['foto_tarihi'] = cells[idx + 1]
                        except (ValueError, IndexError):
                            pass

        self.device_info = device_info
        return device_info

    def extract_all(self, output_dir: str) -> dict:
        """
        Tüm verileri çıkarır.

        Args:
            output_dir: Görüntülerin kaydedileceği klasör

        Returns:
            Tüm çıkarılan verileri içeren sözlük
        """
        self.load()

        images = self.extract_images(output_dir)
        temperature = self.extract_temperature_data()
        device_info = self.extract_device_info()

        # Fotoğraf no: Word içinde FLUKE-xxxxx formatını bul
        import re
        photo_no = ""

        # Tüm tablolarda ve paragraflarda FLUKE-xxxxx pattern'ını ara
        all_text = ""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + cell.text
        for para in self.doc.paragraphs:
            all_text += " " + para.text

        # FLUKE-xxxxx formatını bul (xxxxx = harf ve rakam kombinasyonu)
        fluke_match = re.search(r'FLUKE[-_]?([A-Za-z0-9]+)', all_text, re.IGNORECASE)
        if fluke_match:
            photo_no = fluke_match.group(1)

        # Fallback: dosya adından al
        if not photo_no:
            base = os.path.splitext(os.path.basename(self.docx_path))[0]
            fluke_match = re.search(r'FLUKE[-_]?([A-Za-z0-9]+)', base, re.IGNORECASE)
            if fluke_match:
                photo_no = fluke_match.group(1)
            elif "-" in base:
                photo_no = base.split("-")[-1]

        # Fotoğraf tarihi: Önce Word içinden çek, yoksa cihaz bilgisi veya dosya mtime
        photo_date = device_info.get('foto_tarihi') or ""
        if not photo_date:
            ts = device_info.get('tarih_saat') or ""
            if ts:
                photo_date = ts.split()[0]
            else:
                try:
                    import datetime
                    mtime = os.path.getmtime(self.docx_path)
                    photo_date = datetime.datetime.fromtimestamp(mtime).strftime("%d.%m.%Y")
                except Exception:
                    photo_date = ""

        return {
            'images': images,
            'temperature': temperature,
            'device_info': device_info,
            'photo_no': photo_no,
            'photo_date': photo_date,
        }


# Test için
if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        fluke_path = sys.argv[1]
    else:
        # Relative path or current directory
        fluke_path = "FLUKE-TEST.docx"

    if os.path.exists(fluke_path):
        extractor = FlukeExtractor(fluke_path)

        # Geçici klasöre çıkar
        output_dir = os.path.join(os.path.dirname(os.path.abspath(fluke_path)), "extracted_images")

        result = extractor.extract_all(output_dir)

        print("Cikarilan Goruntuler:")
        for img in result['images']:
            print(f"  - {img}")

        print("\nSicaklik Verileri:")
        for key, value in result['temperature'].items():
            print(f"  {key}: {value}")

        print("\nCihaz Bilgileri:")
        for key, value in result['device_info'].items():
            print(f"  {key}: {value}")
    else:
        print(f"Test dosyasi bulunamadi: {fluke_path}")
